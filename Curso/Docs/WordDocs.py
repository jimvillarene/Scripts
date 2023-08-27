#! /usr/bin/env python3
import docx, os

dir_path = os.path.dirname(os.path.realpath(__file__))
docfile=os.path.join(dir_path,'Demo.docx')
d=docx.Document(docfile)
print(d.paragraphs)
p=d.paragraphs[1]
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[1].bold)
print(p.runs[0].bold==None)
p.runs[3].underline=True
p.runs[3].text='italic and underline'
d.save(docfile)
p.style='Normal'
d.add_paragraph('Hello this is a parragraph')
p=d.paragraphs[3]
p.add_run('This is a new run')
d.save(docfile)
